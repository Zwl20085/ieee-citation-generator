from __future__ import annotations

import argparse
import re
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt
except ModuleNotFoundError as exc:
    raise ModuleNotFoundError(
        "python-docx is required for scripts/export_ieee_docx.py. "
        "Install it with `python -m pip install python-docx`."
    ) from exc


FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
ORDINAL_RE = re.compile(r"(?P<num>\d+)(?P<suffix>st|nd|rd|th)\b")


def normalize_line(line: str) -> str:
    return re.sub(r"\s+", " ", line.replace("\t", " ")).strip()


def split_citations(text: str) -> list[str]:
    return [normalize_line(line) for line in text.splitlines() if line.strip()]


def detect_venue_range(citation: str) -> tuple[int, int] | None:
    quote_candidates = [
        citation.rfind(",”"),
        citation.rfind(',"'),
        citation.rfind(".”"),
        citation.rfind('."'),
        citation.rfind(",?"),
        citation.rfind(".?"),
    ]
    quote_end = max(quote_candidates)
    if quote_end == -1:
        return None

    start = quote_end + 2
    while start < len(citation) and citation[start] == " ":
        start += 1
    if start >= len(citation):
        return None

    remaining = citation[start:]
    if remaining.startswith("in Proc. "):
        venue_start = start + len("in Proc. ")
        boundary = citation.find(", ", venue_start)
        if boundary == -1:
            return None
        return (venue_start, boundary)

    journal_markers = [", vol.", ", no.", ", pp.", ", Art. no.", ", doi:"]
    positions = [citation.find(marker, start) for marker in journal_markers]
    valid = [pos for pos in positions if pos != -1]
    if valid:
        return (start, min(valid))

    website_markers = [". http", ". https", ". www", " (accessed "]
    positions = [citation.find(marker, start) for marker in website_markers]
    valid = [pos for pos in positions if pos != -1]
    if valid:
        return (start, min(valid))

    return None


def add_text_run(paragraph, text: str, italic: bool = False, superscript: bool = False) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = FONT_SIZE
    run.italic = italic
    run.font.superscript = superscript


def add_text_with_ordinals(paragraph, text: str, italic: bool = False) -> None:
    cursor = 0
    for match in ORDINAL_RE.finditer(text):
        add_text_run(paragraph, text[cursor:match.start()], italic=italic)
        add_text_run(paragraph, match.group("num"), italic=italic)
        add_text_run(paragraph, match.group("suffix"), italic=italic, superscript=True)
        cursor = match.end()
    add_text_run(paragraph, text[cursor:], italic=italic)


def add_citation_paragraph(document: Document, citation: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)

    venue_range = detect_venue_range(citation)
    if not venue_range:
        add_text_with_ordinals(paragraph, citation)
        return

    venue_start, venue_end = venue_range
    add_text_with_ordinals(paragraph, citation[:venue_start])
    add_text_with_ordinals(paragraph, citation[venue_start:venue_end], italic=True)
    add_text_with_ordinals(paragraph, citation[venue_end:])


def export_docx(input_path: Path, output_path: Path) -> None:
    citations = split_citations(input_path.read_text(encoding="utf-8"))

    document = Document()
    style = document.styles["Normal"]
    style.font.name = FONT_NAME
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style.font.size = FONT_SIZE

    section = document.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(90)
    section.right_margin = Pt(90)

    for citation in citations:
        add_citation_paragraph(document, citation)

    first = document.paragraphs[0]
    if not first.text and len(document.paragraphs) > len(citations):
        elem = first._element
        elem.getparent().remove(elem)
        first._p = first._element = None

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Export IEEE-formatted citation text from a UTF-8 file into a styled Word document."
    )
    parser.add_argument("input", type=Path, help="Path to a UTF-8 plain-text citation file")
    parser.add_argument(
        "output",
        type=Path,
        nargs="?",
        help="Optional output .docx path. Defaults to <input_basename>_ieee.docx",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    input_path = args.input.resolve()
    if not input_path.is_file():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    output_path = args.output.resolve() if args.output else input_path.with_name(f"{input_path.stem}_ieee.docx")
    export_docx(input_path, output_path)
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
